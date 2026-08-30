"""Format-specific text extraction: .md, .txt, .pdf, .docx.

Each parser takes a file path and returns plain text. Kept deliberately
simple (no layout/table reconstruction) — good enough for Phase 1 chunking.
Swap in a richer parser (e.g. unstructured.io) later if documents have
complex tables/multi-column layouts that this loses.
"""
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx"}


def parse_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def parse_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text)
    return "\n\n".join(pages)


def parse_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def parse_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".md", ".txt"}:
        return parse_text_file(path)
    if suffix == ".pdf":
        return parse_pdf(path)
    if suffix == ".docx":
        return parse_docx(path)
    raise ValueError(f"Unsupported file type: {suffix} ({path})")
